#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
短信供应商管理文档分析工具
基于现有的工时分析框架，扩展支持DOCX和PDF文档处理
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Tuple, Any
import docx
import pdfplumber
import pandas as pd


class DocumentAnalyzer:
    """文档分析器 - 处理DOCX和PDF格式的供应商管理文档"""
    
    def __init__(self, docs_folder: str):
        """
        初始化文档分析器
        
        Args:
            docs_folder: 包含待分析文档的文件夹路径
        """
        self.docs_folder = docs_folder
        self.documents = {}  # 存储提取的文档内容
        self.analysis_results = {}  # 存储分析结果
        self.management_points = []  # 管理要点列表
        
    def extract_docx_content(self, file_path: str) -> str:
        """
        提取DOCX文档的文本内容
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 处理表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n".join(content)
            
        except Exception as e:
            print(f"提取DOCX文档失败 {file_path}: {str(e)}")
            return ""
    
    def extract_pdf_content(self, file_path: str) -> str:
        """
        提取PDF文档的文本内容
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content.append(page_text.strip())
                    
                    # 提取表格内容
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row and any(cell for cell in row if cell and cell.strip()):
                                clean_row = [cell.strip() if cell else "" for cell in row]
                                content.append(" | ".join(clean_row))
            
            return "\n".join(content)
            
        except Exception as e:
            print(f"提取PDF文档失败 {file_path}: {str(e)}")
            return ""
    
    def load_documents(self):
        """加载文件夹中的所有文档"""
        print(f"正在扫描文档文件夹: {self.docs_folder}")
        
        for filename in os.listdir(self.docs_folder):
            file_path = os.path.join(self.docs_folder, filename)
            
            if filename.endswith('.docx'):
                print(f"正在处理DOCX文档: {filename}")
                content = self.extract_docx_content(file_path)
                self.documents[filename] = {
                    'type': 'docx',
                    'content': content,
                    'file_path': file_path
                }
                
            elif filename.endswith('.pdf'):
                print(f"正在处理PDF文档: {filename}")
                content = self.extract_pdf_content(file_path)
                self.documents[filename] = {
                    'type': 'pdf',
                    'content': content,
                    'file_path': file_path
                }
        
        print(f"成功加载 {len(self.documents)} 个文档")
    
    def identify_management_sections(self, content: str) -> Dict[str, List[str]]:
        """
        识别管理办法中的关键章节和条款
        
        Args:
            content: 文档内容
            
        Returns:
            按章节分类的内容字典
        """
        sections = {
            '总则': [],
            '供应商准入': [],
            '供应商评估': [],
            '合同管理': [],
            '服务质量': [],
            '风险管控': [],
            '违约处理': [],
            '其他条款': []
        }
        
        # 关键词匹配模式
        section_patterns = {
            '总则': r'(总则|目的|适用范围|定义)',
            '供应商准入': r'(准入|资质|门槛|条件|要求)',
            '供应商评估': r'(评估|考核|审核|评价|打分)',
            '合同管理': r'(合同|协议|签署|条款)',
            '服务质量': r'(服务质量|SLA|可用性|响应时间|成功率)',
            '风险管控': r'(风险|安全|监控|预警|应急)',
            '违约处理': r'(违约|处罚|终止|解除|责任)',
        }
        
        lines = content.split('\n')
        current_section = '其他条款'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检查是否匹配特定章节
            section_matched = False
            for section, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    current_section = section
                    section_matched = True
                    break
            
            # 将内容添加到对应章节
            if len(line) > 10:  # 过滤过短的行
                sections[current_section].append(line)
        
        return sections
    
    def extract_key_points(self, content: str) -> List[Dict[str, Any]]:
        """
        提取关键管理要点
        
        Args:
            content: 文档内容
            
        Returns:
            关键要点列表
        """
        key_points = []
        
        # 关键要点识别模式
        patterns = {
            '准入标准': r'(准入|资质|门槛).*?([0-9]+.*?年|[0-9]+.*?万|[0-9]+.*?%)',
            '服务指标': r'(可用性|成功率|响应时间).*?([0-9]+.*?%|[0-9]+.*?秒|[0-9]+.*?分钟)',
            '违约条款': r'(违约|处罚|赔偿).*?([0-9]+.*?万|[0-9]+.*?倍|[0-9]+.*?%)',
            '评估周期': r'(评估|考核|审核).*?([0-9]+.*?月|[0-9]+.*?季度|[0-9]+.*?年)',
            '合同期限': r'(合同|协议|期限).*?([0-9]+.*?年|[0-9]+.*?月)',
        }
        
        for point_type, pattern in patterns.items():
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple) and len(match) >= 2:
                    key_points.append({
                        'type': point_type,
                        'description': match[0],
                        'value': match[1],
                        'context': ' '.join(match)
                    })
        
        return key_points
    
    def analyze_content(self, filename: str, content: str) -> Dict[str, Any]:
        """
        分析单个文档内容
        
        Args:
            filename: 文件名
            content: 文档内容
            
        Returns:
            分析结果字典
        """
        # 识别管理章节
        sections = self.identify_management_sections(content)
        
        # 提取关键要点
        key_points = self.extract_key_points(content)
        
        # 统计信息
        stats = {
            'total_lines': len(content.split('\n')),
            'total_chars': len(content),
            'sections_count': len([s for s in sections.values() if s]),
            'key_points_count': len(key_points)
        }
        
        return {
            'filename': filename,
            'sections': sections,
            'key_points': key_points,
            'stats': stats
        }
    
    def compare_documents(self) -> Dict[str, Any]:
        """
        对比分析两份文档的异同
        
        Returns:
            对比分析结果
        """
        if len(self.documents) < 2:
            return {'error': '需要至少两份文档进行对比'}
        
        doc_names = list(self.documents.keys())
        doc1_name, doc2_name = doc_names[0], doc_names[1]
        
        doc1_analysis = self.analysis_results[doc1_name]
        doc2_analysis = self.analysis_results[doc2_name]
        
        comparison = {
            'document_info': {
                'doc1': {'name': doc1_name, 'type': self.documents[doc1_name]['type']},
                'doc2': {'name': doc2_name, 'type': self.documents[doc2_name]['type']}
            },
            'sections_comparison': {},
            'key_points_comparison': {},
            'differences': [],
            'similarities': []
        }
        
        # 对比章节内容
        all_sections = set(doc1_analysis['sections'].keys()) | set(doc2_analysis['sections'].keys())
        for section in all_sections:
            doc1_content = doc1_analysis['sections'].get(section, [])
            doc2_content = doc2_analysis['sections'].get(section, [])
            
            comparison['sections_comparison'][section] = {
                'doc1_items': len(doc1_content),
                'doc2_items': len(doc2_content),
                'doc1_present': len(doc1_content) > 0,
                'doc2_present': len(doc2_content) > 0
            }
        
        # 对比关键要点
        doc1_point_types = set(point['type'] for point in doc1_analysis['key_points'])
        doc2_point_types = set(point['type'] for point in doc2_analysis['key_points'])
        
        comparison['key_points_comparison'] = {
            'doc1_types': list(doc1_point_types),
            'doc2_types': list(doc2_point_types),
            'common_types': list(doc1_point_types & doc2_point_types),
            'doc1_unique': list(doc1_point_types - doc2_point_types),
            'doc2_unique': list(doc2_point_types - doc1_point_types)
        }
        
        return comparison
    
    def generate_management_summary(self) -> List[str]:
        """
        生成供应商管理要点总结
        
        Returns:
            管理要点列表
        """
        summary_points = []
        
        # 基于分析结果生成要点
        all_key_points = []
        for doc_analysis in self.analysis_results.values():
            all_key_points.extend(doc_analysis['key_points'])
        
        # 按类型分组要点
        points_by_type = {}
        for point in all_key_points:
            point_type = point['type']
            if point_type not in points_by_type:
                points_by_type[point_type] = []
            points_by_type[point_type].append(point)
        
        # 生成各类型要点总结
        for point_type, points in points_by_type.items():
            if points:
                summary_points.append(f"**{point_type}**:")
                for point in points[:3]:  # 最多显示3个要点
                    summary_points.append(f"  - {point['context']}")
                summary_points.append("")
        
        return summary_points
    
    def generate_report(self) -> str:
        """
        生成完整的分析报告
        
        Returns:
            Markdown格式的报告内容
        """
        report = []
        
        # 报告标题和时间
        report.append("# 短信供应商管理要点分析报告")
        report.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append("")
        
        # 文档概览
        report.append("## 📋 文档概览")
        report.append(f"**分析文档数量**: {len(self.documents)}")
        for filename, doc_info in self.documents.items():
            report.append(f"- **{filename}** ({doc_info['type'].upper()})")
        report.append("")
        
        # 管理要点总结
        report.append("## 🎯 核心管理要点")
        management_summary = self.generate_management_summary()
        report.extend(management_summary)
        
        # 文档对比分析
        if len(self.documents) >= 2:
            report.append("## 📊 文档对比分析")
            comparison = self.compare_documents()
            
            if 'error' not in comparison:
                doc1_name = comparison['document_info']['doc1']['name']
                doc2_name = comparison['document_info']['doc2']['name']
                
                report.append(f"### 对比文档")
                report.append(f"- **文档1**: {doc1_name}")
                report.append(f"- **文档2**: {doc2_name}")
                report.append("")
                
                # 章节对比
                report.append("### 章节内容对比")
                report.append("| 章节 | 文档1 | 文档2 | 状态 |")
                report.append("|------|-------|-------|------|")
                
                for section, comp in comparison['sections_comparison'].items():
                    doc1_status = "✅" if comp['doc1_present'] else "❌"
                    doc2_status = "✅" if comp['doc2_present'] else "❌"
                    
                    if comp['doc1_present'] and comp['doc2_present']:
                        status = "两者都有"
                    elif comp['doc1_present']:
                        status = "仅文档1有"
                    elif comp['doc2_present']:
                        status = "仅文档2有"
                    else:
                        status = "两者都无"
                    
                    report.append(f"| {section} | {doc1_status} ({comp['doc1_items']}) | {doc2_status} ({comp['doc2_items']}) | {status} |")
                
                report.append("")
                
                # 关键要点对比
                kp_comp = comparison['key_points_comparison']
                if kp_comp['common_types']:
                    report.append("### 共同关注要点")
                    for point_type in kp_comp['common_types']:
                        report.append(f"- {point_type}")
                    report.append("")
                
                if kp_comp['doc1_unique']:
                    report.append(f"### {doc1_name} 独有要点")
                    for point_type in kp_comp['doc1_unique']:
                        report.append(f"- {point_type}")
                    report.append("")
                
                if kp_comp['doc2_unique']:
                    report.append(f"### {doc2_name} 独有要点")
                    for point_type in kp_comp['doc2_unique']:
                        report.append(f"- {point_type}")
                    report.append("")
        
        # 详细分析结果
        report.append("## 📖 详细分析结果")
        for filename, analysis in self.analysis_results.items():
            report.append(f"### {filename}")
            report.append(f"**统计信息**: {analysis['stats']['total_chars']} 字符, {analysis['stats']['sections_count']} 个章节, {analysis['stats']['key_points_count']} 个关键要点")
            report.append("")
            
            # 章节内容
            for section, content_list in analysis['sections'].items():
                if content_list:
                    report.append(f"#### {section}")
                    for content in content_list[:2]:  # 最多显示2条内容
                        report.append(f"- {content[:100]}{'...' if len(content) > 100 else ''}")
                    report.append("")
        
        # 改进建议
        report.append("## 💡 管理改进建议")
        report.append("基于文档分析，建议关注以下方面：")
        report.append("1. **标准化管理流程** - 确保各环节有明确的操作规范")
        report.append("2. **量化考核指标** - 建立可衡量的服务质量标准")
        report.append("3. **风险预警机制** - 完善供应商风险监控体系")
        report.append("4. **定期评估制度** - 建立供应商定期评估和优化机制")
        report.append("")
        
        return "\n".join(report)
    
    def run_full_analysis(self):
        """执行完整的文档分析流程"""
        print("=== 短信供应商管理文档分析 ===")
        
        # 1. 加载文档
        self.load_documents()
        
        if not self.documents:
            print("❌ 未找到可分析的文档文件")
            return
        
        # 2. 分析每个文档
        print("\n正在分析文档内容...")
        for filename, doc_info in self.documents.items():
            print(f"分析文档: {filename}")
            analysis = self.analyze_content(filename, doc_info['content'])
            self.analysis_results[filename] = analysis
        
        # 3. 生成报告
        print("\n正在生成分析报告...")
        report_content = self.generate_report()
        
        # 4. 保存报告
        report_file = '短信供应商管理要点总结.md'
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        print(f"✅ 分析完成！报告已保存为: {report_file}")
        
        # 5. 显示核心要点
        print("\n=== 核心管理要点预览 ===")
        management_summary = self.generate_management_summary()
        for point in management_summary[:10]:  # 显示前10个要点
            print(point)


if __name__ == "__main__":
    # 分析营销短信供应商管理文件夹
    analyzer = DocumentAnalyzer('营销短信供应商管理')
    analyzer.run_full_analysis()